"""
Tests for text extraction from various document sources.

Tests cover:
- Image extraction with OCR
- PDF extraction (text-based and OCR fallback)
- URL/web scraping
- DOCX extraction
- Universal extractor auto-detection
"""

import pytest
import cv2
import numpy as np
from pathlib import Path
from unittest.mock import Mock, patch

from src.extractors import (
    extract_from_image,
    extract_from_pdf,
    extract_from_url,
    extract_from_docx,
    extract_text,
)
from src.ocr import TesseractOCR

# Import private functions for testing
from src import extractors


# ============================================================================
# Fixtures
# ============================================================================

@pytest.fixture
def mock_ocr_engine():
    """Create a mock OCR engine for testing."""
    mock_engine = Mock()
    mock_engine.extract_text.return_value = "Sample extracted text"
    return mock_engine


@pytest.fixture
def sample_image_path(tmp_path: Path) -> Path:
    """Create a simple test image."""
    # Create a white image with some black text-like rectangles
    img = np.ones((100, 200, 3), dtype=np.uint8) * 255
    cv2.rectangle(img, (10, 40), (190, 60), (0, 0, 0), -1)

    image_path = tmp_path / "test_image.png"
    cv2.imwrite(str(image_path), img)
    return image_path


@pytest.fixture
def sample_pdf_path(tmp_path: Path) -> Path:
    """Create a simple test PDF with extractable text."""
    import pypdf
    from pypdf import PdfWriter

    # Create a simple PDF (this is a placeholder - real PDF creation is complex)
    pdf_path = tmp_path / "test.pdf"

    # For testing purposes, we'll create a minimal valid PDF
    # In real tests, you'd use a pre-made test PDF file
    writer = PdfWriter()

    # Note: This creates an empty PDF, but it's enough for testing file existence
    # In a real test suite, you'd want actual test PDFs with content
    with open(pdf_path, 'wb') as f:
        writer.write(f)

    return pdf_path


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> Path:
    """Create a simple test DOCX file."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It has multiple paragraphs.")
    doc.add_paragraph("For testing text extraction.")

    docx_path = tmp_path / "test.docx"
    doc.save(str(docx_path))
    return docx_path


# ============================================================================
# Image Extraction Tests
# ============================================================================

def test_extract_from_image_with_preprocessing(sample_image_path: Path, mock_ocr_engine: Mock):
    """Test image extraction with preprocessing enabled."""
    result = extract_from_image(
        str(sample_image_path),
        mock_ocr_engine,
        preprocess=True,
        denoise_method='gaussian',
        binarize_method='otsu'
    )

    assert isinstance(result, str)
    assert result == "Sample extracted text"
    mock_ocr_engine.extract_text.assert_called_once()


def test_extract_from_image_without_preprocessing(sample_image_path: Path, mock_ocr_engine: Mock):
    """Test image extraction without preprocessing."""
    result = extract_from_image(
        str(sample_image_path),
        mock_ocr_engine,
        preprocess=False
    )

    assert isinstance(result, str)
    assert result == "Sample extracted text"
    mock_ocr_engine.extract_text.assert_called_once()


def test_extract_from_image_file_not_found(mock_ocr_engine: Mock):
    """Test that FileNotFoundError is raised for non-existent image."""
    with pytest.raises(FileNotFoundError, match="Image file not found"):
        extract_from_image("/nonexistent/image.png", mock_ocr_engine)


def test_extract_from_image_invalid_file(tmp_path: Path, mock_ocr_engine: Mock):
    """Test that ValueError is raised for invalid image file."""
    # Create an empty file (not a valid image)
    invalid_file = tmp_path / "invalid.png"
    invalid_file.write_text("not an image")

    with pytest.raises(ValueError, match="Could not load image"):
        extract_from_image(str(invalid_file), mock_ocr_engine, preprocess=False)


# ============================================================================
# PDF Extraction Tests
# ============================================================================

def test_extract_from_pdf_file_not_found(mock_ocr_engine: Mock):
    """Test that FileNotFoundError is raised for non-existent PDF."""
    with pytest.raises(FileNotFoundError, match="PDF file not found"):
        extract_from_pdf("/nonexistent/file.pdf", mock_ocr_engine)


@patch('src.extractors._extract_text_from_pdf')
def test_extract_from_pdf_text_extraction_success(
    mock_extract: Mock,
    sample_pdf_path: Path,
    mock_ocr_engine: Mock
):
    """Test PDF text extraction when PDF has extractable text."""
    mock_extract.return_value = "Extracted PDF text content"

    result = extract_from_pdf(str(sample_pdf_path), mock_ocr_engine, force_ocr=False)

    assert result == "Extracted PDF text content"
    mock_extract.assert_called_once()
    # OCR should not be called since text extraction succeeded
    mock_ocr_engine.extract_text.assert_not_called()


@patch('src.extractors._extract_text_from_pdf')
@patch('src.extractors._ocr_pdf_pages')
def test_extract_from_pdf_fallback_to_ocr(
    mock_ocr_pages: Mock,
    mock_extract: Mock,
    sample_pdf_path: Path,
    mock_ocr_engine: Mock
):
    """Test PDF extraction falls back to OCR when text extraction returns empty."""
    mock_extract.return_value = ""  # Empty text triggers OCR fallback
    mock_ocr_pages.return_value = "OCR extracted text"

    result = extract_from_pdf(str(sample_pdf_path), mock_ocr_engine, force_ocr=False)

    assert result == "OCR extracted text"
    mock_extract.assert_called_once()
    mock_ocr_pages.assert_called_once()


@patch('src.extractors._ocr_pdf_pages')
def test_extract_from_pdf_force_ocr(
    mock_ocr_pages: Mock,
    sample_pdf_path: Path,
    mock_ocr_engine: Mock
):
    """Test PDF extraction with force_ocr=True bypasses text extraction."""
    mock_ocr_pages.return_value = "OCR extracted text"

    result = extract_from_pdf(str(sample_pdf_path), mock_ocr_engine, force_ocr=True)

    assert result == "OCR extracted text"
    mock_ocr_pages.assert_called_once_with(
        str(sample_pdf_path),
        mock_ocr_engine,
        300  # default DPI
    )


def test_extract_text_from_pdf_empty_file(sample_pdf_path: Path):
    """Test _extract_text_from_pdf on an empty PDF."""
    # The sample PDF fixture creates an empty PDF
    result = extractors._extract_text_from_pdf(str(sample_pdf_path))

    # Should return empty string for PDF with no text
    assert isinstance(result, str)
    assert result == ""


# ============================================================================
# URL Extraction Tests
# ============================================================================

@patch('src.extractors.requests.get')
def test_extract_from_url_success(mock_get: Mock):
    """Test successful URL text extraction."""
    # Mock the response
    mock_response = Mock()
    mock_response.content = b"""
    <html>
        <body>
            <article>
                <p>This is a test article with enough content to be extracted.</p>
                <p>It has multiple paragraphs for comprehensive testing.</p>
            </article>
        </body>
    </html>
    """
    mock_response.raise_for_status = Mock()
    mock_get.return_value = mock_response

    result = extract_from_url("https://example.com/article")

    assert isinstance(result, str)
    assert "test article" in result.lower()
    mock_get.assert_called_once()


def test_extract_from_url_invalid_format():
    """Test that ValueError is raised for invalid URL format."""
    with pytest.raises(ValueError, match="Invalid URL format"):
        extract_from_url("not-a-url")


@patch('src.extractors.requests.get')
def test_extract_from_url_with_custom_headers(mock_get: Mock):
    """Test URL extraction with custom user agent."""
    mock_response = Mock()
    mock_response.content = b"<html><body><p>Test content goes here for extraction.</p></body></html>"
    mock_response.raise_for_status = Mock()
    mock_get.return_value = mock_response

    custom_agent = "Custom User Agent"
    extract_from_url("https://example.com", user_agent=custom_agent)

    # Verify custom user agent was used
    call_args = mock_get.call_args
    assert call_args[1]['headers']['User-Agent'] == custom_agent


@patch('src.extractors.requests.get')
def test_extract_from_url_timeout(mock_get: Mock):
    """Test URL extraction with custom timeout."""
    mock_response = Mock()
    mock_response.content = b"<html><body><p>Test content goes here for extraction.</p></body></html>"
    mock_response.raise_for_status = Mock()
    mock_get.return_value = mock_response

    extract_from_url("https://example.com", timeout=60)

    # Verify timeout was passed
    call_args = mock_get.call_args
    assert call_args[1]['timeout'] == 60


def test_extract_article_text_with_article_tag():
    """Test article text extraction finds article tag."""
    from bs4 import BeautifulSoup

    html = """
    <html>
        <body>
            <nav>Navigation should be removed</nav>
            <article>
                <h1>Article Title That Is Long Enough</h1>
                <p>This is the main article content with substantial text.</p>
                <p>Multiple paragraphs for better extraction testing coverage.</p>
            </article>
            <footer>Footer should be removed</footer>
        </body>
    </html>
    """
    soup = BeautifulSoup(html, 'lxml')
    result = extractors._extract_article_text(soup)

    assert "Article Title" in result
    assert "main article content" in result
    assert "Navigation" not in result
    assert "Footer" not in result


def test_extract_article_text_filters_short_snippets():
    """Test that very short text snippets are filtered out."""
    from bs4 import BeautifulSoup

    html = """
    <html>
        <body>
            <p>Short</p>
            <p>This is a longer paragraph that should be included in the output.</p>
        </body>
    </html>
    """
    soup = BeautifulSoup(html, 'lxml')
    result = extractors._extract_article_text(soup)

    assert "Short" not in result
    assert "longer paragraph" in result


# ============================================================================
# TXT Extraction Tests
# ============================================================================

def test_extract_from_txt_success(tmp_path: Path):
    """Test successful TXT text extraction."""
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("This is a test document.\nWith multiple lines.\nFor text extraction.", encoding='utf-8')

    result = extractors.extract_from_txt(str(txt_file))

    assert isinstance(result, str)
    assert "This is a test document." in result
    assert "With multiple lines." in result
    assert "For text extraction." in result


def test_extract_from_txt_file_not_found():
    """Test that FileNotFoundError is raised for non-existent TXT."""
    with pytest.raises(FileNotFoundError, match="TXT file not found"):
        extractors.extract_from_txt("/nonexistent/file.txt")


def test_extract_from_txt_empty_file(tmp_path: Path):
    """Test TXT extraction on empty file."""
    empty_txt = tmp_path / "empty.txt"
    empty_txt.write_text("", encoding='utf-8')

    result = extractors.extract_from_txt(str(empty_txt))
    assert result == ""


def test_extract_from_txt_latin1_encoding(tmp_path: Path):
    """Test TXT extraction with latin-1 encoding fallback."""
    txt_file = tmp_path / "latin1.txt"
    # Write with latin-1 encoding
    txt_file.write_bytes("Café résumé naïve".encode('latin-1'))

    result = extractors.extract_from_txt(str(txt_file))

    assert isinstance(result, str)
    assert len(result) > 0


# ============================================================================
# DOCX Extraction Tests
# ============================================================================

def test_extract_from_docx_success(sample_docx_path: Path):
    """Test successful DOCX text extraction."""
    result = extract_from_docx(str(sample_docx_path))

    assert isinstance(result, str)
    assert "test document" in result
    assert "multiple paragraphs" in result
    assert "text extraction" in result


def test_extract_from_docx_file_not_found():
    """Test that FileNotFoundError is raised for non-existent DOCX."""
    with pytest.raises(FileNotFoundError, match="DOCX file not found"):
        extract_from_docx("/nonexistent/file.docx")


def test_extract_from_docx_empty_document(tmp_path: Path):
    """Test DOCX extraction on empty document."""
    from docx import Document

    doc = Document()
    empty_docx = tmp_path / "empty.docx"
    doc.save(str(empty_docx))

    result = extract_from_docx(str(empty_docx))

    assert isinstance(result, str)
    assert result == ""


# ============================================================================
# Universal Extractor Tests
# ============================================================================

def test_detect_source_type_url():
    """Test source type detection for URLs."""
    assert extractors._detect_source_type("http://example.com") == "url"
    assert extractors._detect_source_type("https://example.com") == "url"


def test_detect_source_type_image():
    """Test source type detection for image files."""
    assert extractors._detect_source_type("document.jpg") == "image"
    assert extractors._detect_source_type("photo.PNG") == "image"
    assert extractors._detect_source_type("scan.tiff") == "image"


def test_detect_source_type_pdf():
    """Test source type detection for PDF files."""
    assert extractors._detect_source_type("document.pdf") == "pdf"
    assert extractors._detect_source_type("file.PDF") == "pdf"


def test_detect_source_type_docx():
    """Test source type detection for DOCX files."""
    assert extractors._detect_source_type("document.docx") == "docx"
    assert extractors._detect_source_type("file.DOCX") == "docx"


def test_detect_source_type_txt():
    """Test source type detection for TXT files."""
    assert extractors._detect_source_type("document.txt") == "txt"
    assert extractors._detect_source_type("file.TXT") == "txt"


def test_detect_source_type_unknown():
    """Test that ValueError is raised for unknown file types."""
    with pytest.raises(ValueError, match="Cannot detect source type"):
        extractors._detect_source_type("file.xyz")


@patch('src.extractors.extract_from_image')
def test_extract_text_auto_detect_image(mock_extract: Mock, sample_image_path: Path):
    """Test universal extractor auto-detects and extracts from image."""
    mock_extract.return_value = "Image text"

    result = extract_text(str(sample_image_path))

    assert result == "Image text"
    mock_extract.assert_called_once()


@patch('src.extractors.extract_from_pdf')
def test_extract_text_auto_detect_pdf(mock_extract: Mock, sample_pdf_path: Path):
    """Test universal extractor auto-detects and extracts from PDF."""
    mock_extract.return_value = "PDF text"

    result = extract_text(str(sample_pdf_path))

    assert result == "PDF text"
    mock_extract.assert_called_once()


@patch('src.extractors.extract_from_url')
def test_extract_text_auto_detect_url(mock_extract: Mock):
    """Test universal extractor auto-detects and extracts from URL."""
    mock_extract.return_value = "URL text"

    result = extract_text("https://example.com/article")

    assert result == "URL text"
    mock_extract.assert_called_once()


@patch('src.extractors.extract_from_docx')
def test_extract_text_auto_detect_docx(mock_extract: Mock, sample_docx_path: Path):
    """Test universal extractor auto-detects and extracts from DOCX."""
    mock_extract.return_value = "DOCX text"

    result = extract_text(str(sample_docx_path))

    assert result == "DOCX text"
    mock_extract.assert_called_once()


@patch('src.extractors.extract_from_txt')
def test_extract_text_auto_detect_txt(mock_extract: Mock, tmp_path: Path):
    """Test universal extractor auto-detects and extracts from TXT."""
    mock_extract.return_value = "TXT text"

    txt_file = tmp_path / "test.txt"
    txt_file.write_text("test content", encoding='utf-8')

    result = extract_text(str(txt_file))

    assert result == "TXT text"
    mock_extract.assert_called_once()


@patch('src.extractors.extract_from_image')
def test_extract_text_force_source_type(mock_extract: Mock, sample_pdf_path: Path):
    """Test universal extractor with forced source type."""
    mock_extract.return_value = "Forced extraction"

    # Force treating PDF as image
    result = extract_text(str(sample_pdf_path), source_type='image')

    assert result == "Forced extraction"
    mock_extract.assert_called_once()


def test_extract_text_uses_default_ocr_engine(sample_image_path: Path):
    """Test that extract_text creates default OCR engine when none provided."""
    with patch('src.extractors.extract_from_image') as mock_extract:
        mock_extract.return_value = "Text"

        extract_text(str(sample_image_path))

        # Should have called extract_from_image with a TesseractOCR instance
        mock_extract.assert_called_once()
        call_args = mock_extract.call_args
        ocr_engine = call_args[0][1]
        assert isinstance(ocr_engine, TesseractOCR)


def test_extract_text_invalid_source_type():
    """Test that ValueError is raised for invalid source type."""
    with pytest.raises(ValueError, match="Unsupported source type"):
        extract_text("file.pdf", source_type="invalid")


# ============================================================================
# Integration Tests
# ============================================================================

@pytest.mark.slow
def test_extract_from_image_real_ocr(sample_image_path: Path):
    """Integration test with real Tesseract OCR."""
    ocr_engine = TesseractOCR()

    result = extract_from_image(str(sample_image_path), ocr_engine, preprocess=True)

    # Should return a string (might be empty for our simple test image)
    assert isinstance(result, str)


@pytest.mark.slow
def test_extract_text_real_integration(sample_image_path: Path):
    """Integration test of universal extractor with real OCR."""
    result = extract_text(str(sample_image_path))

    assert isinstance(result, str)
